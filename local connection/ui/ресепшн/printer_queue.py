import docx
import datetime
from docx import Document
from docx.shared import Pt

class printer(object):
	def print_queue(text):
		doc = docx.Document('Очередь Шаблон.docx')
		all_paras = doc.paragraphs
		style = doc.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		para1 = all_paras[2]
		para1.add_run("123").font.size = docx.shared.Pt(48)
		para2 = all_paras[4]
		para2.add_run(text).font.size = docx.shared.Pt(11)
		doc.save("Очередь Шаблон1.docx")




# import docx
# import datetime
# from docx import Document
# from docx.shared import Pt

# kyzmet_name = "Pezda"

# doc = docx.Document('Очередь Шаблон.docx')
# all_paras = doc.paragraphs
# style = doc.styles['Normal']
# font = style.font
# font.name = 'Times New Roman'
# para1 = all_paras[2]
# para1.add_run("123").font.size = docx.shared.Pt(48)
# para2 = all_paras[4]
# para2.add_run(kyzmet_name).font.size = docx.shared.Pt(11)
# doc.save("Очередь Шаблон1.docx")