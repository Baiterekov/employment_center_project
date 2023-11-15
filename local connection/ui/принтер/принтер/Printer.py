# import os
import docx

doc = docx.Document("Очередь Шаблон.docx")
all_paras = doc.paragraphs
para1 = all_paras[2]
para1.add_run("125")
doc.save("Очередь Шаблон.docx")
for para in all_paras:
    print(para.text)
    print("-------")
# os.startfile("Очередь Шаблон.docx", "print")
