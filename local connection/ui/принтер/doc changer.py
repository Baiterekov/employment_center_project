# import docx

# doc = docx.Document('example.docx')


# all_paras = doc.paragraphs
# print(len(all_paras))

# queueue_number = all_paras[0]
# queueue_number.text = None
# queueue_number.add_run("15")

# doc.save("example.docx")

# import os
import docx
from docx import Document
from docx.shared import Pt


doc = docx.Document('Очередь Шаблон.docx')
all_paras = doc.paragraphs
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
para1 = all_paras[2]
para1.add_run("5").font.size = docx.shared.Pt(48)

doc.save("Очередь Шаблон1.docx")